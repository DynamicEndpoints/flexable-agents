"""Document processing and analysis tools for MCP server."""

import logging
import io
from typing import Dict, Any, List, Optional
from pathlib import Path
from datetime import datetime
import json
import re
import time

# Document processing libraries
try:
    import PyPDF2
    import docx
    from PIL import Image
    import pytesseract
    DOCUMENT_LIBS_AVAILABLE = True
except ImportError:
    DOCUMENT_LIBS_AVAILABLE = False

from src.core import mcp_tool
from src.mcp.logging_system import log_request_metrics
from src.tools.m365_tools import with_error_handling

logger = logging.getLogger(__name__)


def register_document_tools(server, config):
    """Register document processing tools with the MCP server."""
    
    if not DOCUMENT_LIBS_AVAILABLE:
        logger.warning("Some document processing libraries not available. Install PyPDF2, python-docx, PIL, pytesseract for full functionality.")
    
    server.register_tool(
        name="document_processing",
        description="Process and analyze documents (PDF, Word, etc.)",
        handler=document_processing,
        returns="Document processing result"
    )
    
    server.register_tool(
        name="document_conversion",
        description="Convert documents between formats",
        handler=document_conversion,
        returns="Document conversion result"
    )
    
    server.register_tool(
        name="document_analysis",
        description="Analyze document content and structure",
        handler=document_analysis,
        returns="Document analysis result"
    )
    
    server.register_tool(
        name="ocr_processing",
        description="Extract text from images and scanned documents using OCR",
        handler=ocr_processing,
        returns="OCR processing result"
    )
    
    logger.info("Registered document processing tools")


@mcp_tool(
    name="document_processing",
    description="Process and analyze documents",
    category="documents",
    tags=["documents", "text", "processing"]
)
@with_error_handling("document_processing")
async def document_processing(
    file_path: str, 
    operation: str = "extract_text", 
    options: Optional[Dict[str, Any]] = None
) -> Dict[str, Any]:
    """
    Process and analyze documents.
    
    Args:
        file_path: Path to the document file
        operation: Operation to perform (extract_text, extract_metadata, count_words, find_patterns)
        options: Additional options for processing
    """
    try:
        options = options or {}
        file_path = Path(file_path)
        
        if not file_path.exists():
            return {"status": "error", "message": f"File not found: {file_path}"}
        
        file_extension = file_path.suffix.lower()
        
        if operation == "extract_text":
            if file_extension == ".pdf":
                text = extract_pdf_text(file_path)
            elif file_extension in [".docx", ".doc"]:
                text = extract_word_text(file_path)
            elif file_extension == ".txt":
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
            else:
                return {"status": "error", "message": f"Unsupported file format: {file_extension}"}
            
            return {
                "status": "success",
                "message": "Text extracted successfully",
                "data": {
                    "file_path": str(file_path),
                    "text": text,
                    "character_count": len(text),
                    "word_count": len(text.split()),
                    "file_type": file_extension
                }
            }
            
        elif operation == "extract_metadata":
            metadata = {
                "file_name": file_path.name,
                "file_size": file_path.stat().st_size,
                "file_extension": file_extension,
                "created_date": datetime.fromtimestamp(file_path.stat().st_ctime).isoformat(),
                "modified_date": datetime.fromtimestamp(file_path.stat().st_mtime).isoformat()
            }
            
            # Add format-specific metadata
            if file_extension == ".pdf":
                metadata.update(extract_pdf_metadata(file_path))
            elif file_extension in [".docx", ".doc"]:
                metadata.update(extract_word_metadata(file_path))
            
            return {
                "status": "success",
                "message": "Metadata extracted successfully",
                "data": metadata
            }
            
        elif operation == "count_words":
            if file_extension == ".pdf":
                text = extract_pdf_text(file_path)
            elif file_extension in [".docx", ".doc"]:
                text = extract_word_text(file_path)
            elif file_extension == ".txt":
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
            else:
                return {"status": "error", "message": f"Unsupported file format: {file_extension}"}
            
            words = text.split()
            unique_words = set(word.lower().strip('.,!?;:"()[]{}') for word in words)
            
            return {
                "status": "success",
                "message": "Word count completed",
                "data": {
                    "total_words": len(words),
                    "unique_words": len(unique_words),
                    "characters": len(text),
                    "characters_no_spaces": len(text.replace(' ', '')),
                    "paragraphs": len(text.split('\n\n')),
                    "lines": len(text.split('\n'))
                }
            }
            
        elif operation == "find_patterns":
            pattern = options.get("pattern")
            if not pattern:
                return {"status": "error", "message": "Pattern required for find_patterns operation"}
            
            if file_extension == ".pdf":
                text = extract_pdf_text(file_path)
            elif file_extension in [".docx", ".doc"]:
                text = extract_word_text(file_path)
            elif file_extension == ".txt":
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
            else:
                return {"status": "error", "message": f"Unsupported file format: {file_extension}"}
            
            matches = re.findall(pattern, text, re.IGNORECASE if options.get("case_insensitive") else 0)
            
            return {
                "status": "success",
                "message": f"Found {len(matches)} matches for pattern",
                "data": {
                    "pattern": pattern,
                    "matches": matches,
                    "match_count": len(matches),
                    "unique_matches": list(set(matches))
                }
            }
            
        else:
            return {"status": "error", "message": f"Unknown operation: {operation}"}
            
    except Exception as e:
        logger.error(f"Error in document processing: {e}")
        return {"status": "error", "message": str(e)}


@mcp_tool(
    name="document_conversion",
    description="Convert documents between formats",
    category="documents",
    tags=["conversion", "formats", "transformation"]
)
@with_error_handling("document_conversion")
async def document_conversion(
    input_file: str,
    output_file: str,
    conversion_type: str,
    options: Optional[Dict[str, Any]] = None
) -> Dict[str, Any]:
    """
    Convert documents between different formats.
    
    Args:
        input_file: Input file path
        output_file: Output file path
        conversion_type: Type of conversion (pdf_to_text, word_to_text, text_to_html)
        options: Conversion options
    """
    try:
        options = options or {}
        input_path = Path(input_file)
        output_path = Path(output_file)
        
        if not input_path.exists():
            return {"status": "error", "message": f"Input file not found: {input_file}"}
        
        if conversion_type == "pdf_to_text":
            text = extract_pdf_text(input_path)
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text)
                
        elif conversion_type == "word_to_text":
            text = extract_word_text(input_path)
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text)
                
        elif conversion_type == "text_to_html":
            with open(input_path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            # Simple text to HTML conversion
            html_content = f"""<!DOCTYPE html>
<html>
<head>
    <title>{input_path.stem}</title>
    <meta charset="utf-8">
</head>
<body>
    <h1>{input_path.stem}</h1>
    <pre>{text}</pre>
</body>
</html>"""
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
                
        else:
            return {"status": "error", "message": f"Unknown conversion type: {conversion_type}"}
        
        return {
            "status": "success",
            "message": f"Document converted successfully: {conversion_type}",
            "data": {
                "input_file": str(input_path),
                "output_file": str(output_path),
                "conversion_type": conversion_type,
                "output_size": output_path.stat().st_size
            }
        }
        
    except Exception as e:
        logger.error(f"Error in document conversion: {e}")
        return {"status": "error", "message": str(e)}


@mcp_tool(
    name="document_analysis",
    description="Analyze document content and structure",
    category="documents",
    tags=["analysis", "content", "structure"]
)
@with_error_handling("document_analysis")
async def document_analysis(
    file_path: str,
    analysis_type: str = "content",
    options: Optional[Dict[str, Any]] = None
) -> Dict[str, Any]:
    """
    Analyze document content and structure.
    
    Args:
        file_path: Path to the document file
        analysis_type: Type of analysis (content, structure, readability, keywords)
        options: Analysis options
    """
    try:
        options = options or {}
        file_path = Path(file_path)
        
        if not file_path.exists():
            return {"status": "error", "message": f"File not found: {file_path}"}
        
        file_extension = file_path.suffix.lower()
        
        # Extract text based on file type
        if file_extension == ".pdf":
            text = extract_pdf_text(file_path)
        elif file_extension in [".docx", ".doc"]:
            text = extract_word_text(file_path)
        elif file_extension == ".txt":
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
        else:
            return {"status": "error", "message": f"Unsupported file format: {file_extension}"}
        
        if analysis_type == "content":
            sentences = text.split('.')
            paragraphs = text.split('\n\n')
            words = text.split()
            
            analysis_result = {
                "word_count": len(words),
                "sentence_count": len([s for s in sentences if s.strip()]),
                "paragraph_count": len([p for p in paragraphs if p.strip()]),
                "character_count": len(text),
                "average_words_per_sentence": len(words) / max(len(sentences), 1),
                "average_sentences_per_paragraph": len(sentences) / max(len(paragraphs), 1)
            }
            
        elif analysis_type == "keywords":
            words = re.findall(r'\b\w+\b', text.lower())
            word_freq = {}
            for word in words:
                if len(word) > 3:  # Only consider words longer than 3 characters
                    word_freq[word] = word_freq.get(word, 0) + 1
            
            # Get top keywords
            top_keywords = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)[:20]
            
            analysis_result = {
                "total_words": len(words),
                "unique_words": len(word_freq),
                "top_keywords": dict(top_keywords),
                "vocabulary_richness": len(word_freq) / max(len(words), 1)
            }
            
        elif analysis_type == "structure":
            lines = text.split('\n')
            
            # Identify potential headings (lines that are shorter and followed by longer lines)
            potential_headings = []
            for i, line in enumerate(lines):
                if line.strip() and len(line.strip()) < 100:
                    if i + 1 < len(lines) and len(lines[i + 1].strip()) > len(line.strip()):
                        potential_headings.append(line.strip())
            
            analysis_result = {
                "total_lines": len(lines),
                "non_empty_lines": len([line for line in lines if line.strip()]),
                "potential_headings": potential_headings,
                "structure_analysis": {
                    "has_title": len(potential_headings) > 0,
                    "heading_count": len(potential_headings),
                    "average_line_length": sum(len(line) for line in lines) / max(len(lines), 1)
                }
            }
            
        else:
            return {"status": "error", "message": f"Unknown analysis type: {analysis_type}"}
        
        return {
            "status": "success",
            "message": f"Document analysis completed: {analysis_type}",
            "data": {
                "file_path": str(file_path),
                "analysis_type": analysis_type,
                "timestamp": datetime.now().isoformat(),
                "results": analysis_result
            }
        }
        
    except Exception as e:
        logger.error(f"Error in document analysis: {e}")
        return {"status": "error", "message": str(e)}


@mcp_tool(
    name="ocr_processing",
    description="Extract text from images using OCR",
    category="documents",
    tags=["ocr", "images", "text extraction"]
)
@with_error_handling("ocr_processing")
async def ocr_processing(
    image_path: str,
    options: Optional[Dict[str, Any]] = None
) -> Dict[str, Any]:
    """
    Extract text from images and scanned documents using OCR.
    
    Args:
        image_path: Path to the image file
        options: OCR processing options
    """
    try:
        options = options or {}
        image_path = Path(image_path)
        
        if not image_path.exists():
            return {"status": "error", "message": f"Image file not found: {image_path}"}
        
        try:
            # Use pytesseract for OCR
            image = Image.open(image_path)
            extracted_text = pytesseract.image_to_string(image)
            
            # Get additional OCR data
            ocr_data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)
            
            # Calculate confidence scores
            confidences = [int(conf) for conf in ocr_data['conf'] if int(conf) > 0]
            average_confidence = sum(confidences) / len(confidences) if confidences else 0
            
            return {
                "status": "success",
                "message": "OCR processing completed",
                "data": {
                    "image_path": str(image_path),
                    "extracted_text": extracted_text,
                    "word_count": len(extracted_text.split()),
                    "character_count": len(extracted_text),
                    "average_confidence": round(average_confidence, 2),
                    "image_dimensions": image.size,
                    "processing_timestamp": datetime.now().isoformat()
                }
            }
            
        except Exception as ocr_error:
            return {
                "status": "error",
                "message": f"OCR processing failed: {str(ocr_error)}. Ensure pytesseract is installed and configured."
            }
        
    except Exception as e:
        logger.error(f"Error in OCR processing: {e}")
        return {"status": "error", "message": str(e)}


# Helper functions
def extract_pdf_text(file_path: Path) -> str:
    """Extract text from PDF file."""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text
    except Exception as e:
        logger.error(f"Error extracting PDF text: {e}")
        return f"Error extracting PDF text: {str(e)}"


def extract_word_text(file_path: Path) -> str:
    """Extract text from Word document."""
    try:
        doc = docx.Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        logger.error(f"Error extracting Word text: {e}")
        return f"Error extracting Word text: {str(e)}"


def extract_pdf_metadata(file_path: Path) -> dict:
    """Extract metadata from PDF file."""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            metadata = pdf_reader.metadata
            return {
                "title": metadata.get('/Title', ''),
                "author": metadata.get('/Author', ''),
                "subject": metadata.get('/Subject', ''),
                "creator": metadata.get('/Creator', ''),
                "producer": metadata.get('/Producer', ''),
                "creation_date": str(metadata.get('/CreationDate', '')),
                "modification_date": str(metadata.get('/ModDate', '')),
                "page_count": len(pdf_reader.pages)
            }
    except Exception as e:
        logger.error(f"Error extracting PDF metadata: {e}")
        return {"error": str(e)}


def extract_word_metadata(file_path: Path) -> dict:
    """Extract metadata from Word document."""
    try:
        doc = docx.Document(file_path)
        core_props = doc.core_properties
        return {
            "title": core_props.title or '',
            "author": core_props.author or '',
            "subject": core_props.subject or '',
            "keywords": core_props.keywords or '',
            "comments": core_props.comments or '',
            "created": str(core_props.created) if core_props.created else '',
            "modified": str(core_props.modified) if core_props.modified else '',
            "last_modified_by": core_props.last_modified_by or '',
            "paragraph_count": len(doc.paragraphs)
        }
    except Exception as e:
        logger.error(f"Error extracting Word metadata: {e}")
        return {"error": str(e)}
